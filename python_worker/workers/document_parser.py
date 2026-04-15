"""
Document Parser Worker for Harbor Platform

Extracts text from student-uploaded documents (PDF and DOCX) stored in
Supabase Storage, then feeds the text to Gemini AI for resume analysis.

Supported formats:
  - PDF  — extracted via PyMuPDF (fitz), handles text-based PDFs
  - DOCX — extracted via python-docx (paragraphs + table cells)

Page limits and size caps prevent runaway processing on huge files.
"""

import logging
from typing import Optional
import fitz           # PyMuPDF
from docx import Document
import io

from db import supabase
from llm import analyze_resume_with_ai

logger = logging.getLogger(__name__)

_MAX_PAGES = 20          # Hard cap: skip pages beyond this for very large PDFs
_MAX_CHARS = 10_000      # Send at most this many chars to Gemini (token safety)


async def process_pdf(bucket_name: str, file_path: str, file_url: str) -> str:
    """
    Downloads a PDF from Supabase Storage and extracts its text content.

    Args:
        bucket_name: Supabase storage bucket (e.g. "resumes").
        file_path:   Path inside the bucket.
        file_url:    Public URL (used only in log messages).

    Returns:
        Extracted text string, or an "Error: ..." message on failure.
    """
    logger.info(f"Processing PDF — bucket: {bucket_name}, path: {file_path}")

    if not (file_path and bucket_name and supabase):
        logger.warning("Missing bucket/path or Supabase not initialized")
        return "Error: Unable to download document — missing configuration"

    try:
        file_bytes = supabase.storage.from_(bucket_name).download(file_path)
        logger.info(f"Downloaded {len(file_bytes):,} bytes from Supabase")
    except Exception as e:
        logger.error(f"Supabase download failed: {e}")
        raise

    if not file_bytes:
        raise ValueError("Downloaded file is empty")

    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
    except fitz.FileDataError as e:
        raise Exception(f"PDF file is corrupted or invalid: {e}")

    total_pages = len(pdf)
    pages_to_read = min(total_pages, _MAX_PAGES)
    logger.info(f"PDF opened: {total_pages} pages (reading up to {pages_to_read})")

    text_parts = []
    for page_num in range(pages_to_read):
        page_text = pdf.load_page(page_num).get_text()
        if page_text.strip():
            text_parts.append(page_text)

    pdf.close()

    extracted = "\n".join(text_parts).strip()

    if not extracted:
        return (
            "Error: PDF contains no extractable text. "
            "Ensure your resume is text-based, not a scanned image."
        )

    if total_pages > _MAX_PAGES:
        logger.warning(f"PDF has {total_pages} pages — only first {_MAX_PAGES} processed")

    logger.info(f"Extracted {len(extracted):,} characters from PDF")
    return extracted


async def process_docx(bucket_name: str, file_path: str) -> str:
    """
    Downloads a DOCX from Supabase Storage and extracts its text.
    Reads all paragraphs and table cell content.

    Args:
        bucket_name: Supabase storage bucket.
        file_path:   Path inside the bucket.

    Returns:
        Extracted text string, or an "Error: ..." message on failure.
    """
    logger.info(f"Processing DOCX — bucket: {bucket_name}, path: {file_path}")

    if not (file_path and bucket_name and supabase):
        return "Error: Unable to download document — missing configuration"

    try:
        file_bytes = supabase.storage.from_(bucket_name).download(file_path)
        logger.info(f"Downloaded {len(file_bytes):,} bytes from Supabase")
    except Exception as e:
        logger.error(f"Supabase download failed: {e}")
        raise

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise Exception(f"DOCX file is corrupted or invalid: {e}")

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    extracted = "\n".join(text_parts).strip()

    if not extracted:
        return "Error: DOCX contains no extractable text."

    logger.info(f"Extracted {len(extracted):,} characters from DOCX")
    return extracted


async def extract_document_text(
    bucket_name: str,
    file_path: str,
    file_url: str,
    document_type: str = "resume",
) -> str:
    """
    Dispatcher — routes to PDF or DOCX extractor based on file extension.

    Args:
        bucket_name:   Supabase storage bucket.
        file_path:     Path inside the bucket (used to detect extension).
        file_url:      Public URL (passed to PDF handler for logging).
        document_type: Type hint ("resume" or "document").

    Returns:
        Extracted text string, or "Error: ..." on failure.
    """
    lower_path = (file_path or "").lower()

    if lower_path.endswith(".docx"):
        return await process_docx(bucket_name, file_path)
    else:
        # Default: treat as PDF
        return await process_pdf(bucket_name, file_path, file_url)
